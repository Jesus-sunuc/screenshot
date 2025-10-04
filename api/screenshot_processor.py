import logging
from typing import List, Dict

import cv2
import numpy as np
import pytesseract
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Inches, RGBColor

logger = logging.getLogger(__name__)


class ScreenshotProcessor:
    CONFIDENCE_THRESHOLD = 30
    SUPPORTED_FORMATS = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp'}

    def __init__(self):
        self.font_size_thresholds = {
            'heading1': 22,
            'heading2': 18,
            'heading3': 14,
            'body': 11
        }

    def process_image(self, image_path: str, output_path: str) -> None:
        logger.info(f"Processing: {image_path}")

        img = cv2.imread(image_path)
        if img is None:
            raise ValueError(f"Could not read image: {image_path}")

        text_blocks = self.extract_text_with_layout(img)
        logger.info(f"Extracted {len(text_blocks)} blocks")

        self.create_word_document(text_blocks, output_path)
        logger.info(f"Created: {output_path}")

    def extract_text_with_layout(self, img: np.ndarray) -> List[Dict]:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        data = pytesseract.image_to_data(gray, output_type=pytesseract.Output.DICT)

        text_blocks = []
        current_line = self._create_empty_line()

        for i in range(len(data['text'])):
            if int(data['conf'][i]) > self.CONFIDENCE_THRESHOLD:
                text = data['text'][i].strip()
                if text:
                    if self._is_new_line(current_line, data, i):
                        text_blocks.append(current_line.copy())
                        current_line = self._create_line_from_data(data, i, text)
                    else:
                        self._append_to_current_line(current_line, data, i, text)

        if current_line['text']:
            text_blocks.append(current_line)

        for block in text_blocks:
            block['estimated_font_size'] = self.estimate_font_size(block['height'])
            block['style'] = self.determine_style(block)

        return text_blocks

    def _create_empty_line(self) -> Dict:
        return {
            'text': '',
            'left': 0,
            'top': 0,
            'height': 0,
            'confidence': 0,
            'block_num': 0
        }

    def _is_new_line(self, current_line: Dict, data: Dict, index: int) -> bool:
        return (current_line['text'] and
                (data['block_num'][index] != current_line['block_num'] or
                 abs(data['top'][index] - current_line['top']) > 10))

    def _create_line_from_data(self, data: Dict, index: int, text: str) -> Dict:
        return {
            'text': text,
            'left': data['left'][index],
            'top': data['top'][index],
            'height': data['height'][index],
            'confidence': data['conf'][index],
            'block_num': data['block_num'][index]
        }

    def _append_to_current_line(self, current_line: Dict, data: Dict, index: int, text: str) -> None:
        if current_line['text']:
            current_line['text'] += ' ' + text
        else:
            current_line['text'] = text
            current_line['left'] = data['left'][index]
            current_line['top'] = data['top'][index]
            current_line['height'] = data['height'][index]
            current_line['block_num'] = data['block_num'][index]

    def estimate_font_size(self, height: int) -> int:
        return max(8, int(height * 0.8))

    def determine_style(self, block: Dict) -> str:
        font_size = block['estimated_font_size']
        text = block['text']

        is_all_caps = text.isupper() and len(text.split()) <= 15
        is_short_standalone = len(text.split()) <= 6 and font_size >= 10

        if is_all_caps and font_size < 14:
            return 'header'
        elif font_size >= self.font_size_thresholds['heading1']:
            return 'heading1'
        elif font_size >= self.font_size_thresholds['heading2'] or is_short_standalone:
            return 'heading2'
        elif font_size >= self.font_size_thresholds['heading3']:
            return 'heading3'
        else:
            return 'body'

    def create_word_document(self, text_blocks: List[Dict], output_path: str) -> None:
        doc = Document()
        self._configure_document_style(doc)

        last_top = 0
        last_block_num = -1

        for i, block in enumerate(text_blocks):
            self._add_paragraph_to_document(doc, block, i, last_top, last_block_num)
            last_top = block['top'] + block['height']
            last_block_num = block['block_num']

        doc.save(output_path)

    def _configure_document_style(self, doc: Document) -> None:
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

    def _add_paragraph_to_document(self, doc: Document, block: Dict, index: int,
                                   last_top: int, last_block_num: int) -> None:
        text = block['text']
        style_type = block['style']

        spacing_pixels = block['top'] - last_top if last_top > 0 else 0
        spacing_before = max(0, int(spacing_pixels * 0.3))
        is_new_paragraph = spacing_pixels > 20 or block['block_num'] != last_block_num

        if style_type == 'header':
            p = doc.add_paragraph(text)
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
            p.runs[0].font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)

        elif style_type == 'heading1':
            p = doc.add_heading(text, level=1)
            p.runs[0].font.size = Pt(block['estimated_font_size'])
            p.runs[0].font.bold = True
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_before = Pt(spacing_before)
            p.paragraph_format.space_after = Pt(8)

        elif style_type == 'heading2':
            p = doc.add_heading(text, level=2)
            p.runs[0].font.size = Pt(block['estimated_font_size'])
            p.runs[0].font.bold = True
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_before = Pt(max(12, spacing_before))
            p.paragraph_format.space_after = Pt(6)

        elif style_type == 'heading3':
            p = doc.add_heading(text, level=3)
            p.runs[0].font.size = Pt(block['estimated_font_size'])
            p.runs[0].font.bold = True
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_before = Pt(max(10, spacing_before))
            p.paragraph_format.space_after = Pt(4)

        else:
            p = doc.add_paragraph(text)
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

            if is_new_paragraph and index > 0:
                p.paragraph_format.space_before = Pt(max(8, spacing_before))
            else:
                p.paragraph_format.space_before = Pt(0)

            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.15

        if style_type != 'header':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
